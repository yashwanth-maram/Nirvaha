from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title_run = title.add_run('Nirvaha Website Structure')
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(31, 78, 121)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()  # Spacing

# Document metadata
metadata = doc.add_paragraph('Internal Documentation | Structure Overview')
metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
metadata_run = metadata.runs[0]
metadata_run.font.size = Pt(10)
metadata_run.font.italic = True
doc.add_paragraph()  # Spacing

# Section 1: Introduction
doc.add_heading('1. Introduction', level=1)
intro_text = """This document provides a comprehensive overview of the Nirvaha website structure and organization. It outlines the complete layout, navigation flow, key sections, and core platform features that constitute the Nirvaha digital experience. This document is intended for internal review and serves as a reference guide for understanding how the website is organized and how users interact with its various components."""
doc.add_paragraph(intro_text)
doc.add_paragraph()

# Section 2: Overall Website Flow
doc.add_heading('2. Overall Website Flow', level=1)
flow_text = """The Nirvaha website follows a clear top-to-bottom navigation structure, beginning with an engaging landing page and progressing through various sections that introduce users to the platform's offerings. The flow is designed to guide visitors from initial awareness of Nirvaha's mission and services, through platform features and community engagement, to eventual conversion and participation. The structure ensures that users can easily navigate to key features, understand the value proposition, and access authentication and registration flows at appropriate touchpoints throughout the experience."""
doc.add_paragraph(flow_text)
doc.add_paragraph()

# Section 3: Landing Page Structure
doc.add_heading('3. Landing Page Structure', level=1)
landing_intro = """The landing page is the primary entry point for all users and is structured to guide them through the Nirvaha experience in a logical sequence. The following sections appear in exact order:"""
doc.add_paragraph(landing_intro)

landing_sections = [
    "Hero Section - The primary visual introduction to Nirvaha",
    "About Us Section - Overview of Nirvaha's mission and vision",
    "Services Section - High-level introduction to core services offered",
    "Meditation & Wellness Features - Highlighting meditation and wellness capabilities",
    "Join Our Healing Community - Call-to-action for community participation",
    "User Feedback / Testimonials - Social proof through user testimonials",
    "Gallery - Visual showcase of platform features and offerings",
    "Collaborators / Partners - Display of organizational partnerships",
    "Contact Form - Direct communication channel for inquiries",
    "Footer - Site-wide navigation and legal information"
]

for section in landing_sections:
    doc.add_paragraph(section, style='List Bullet')
doc.add_paragraph()

# Section 4: Section-wise Breakdown
doc.add_heading('4. Section-wise Breakdown', level=1)

sections_breakdown = [
    {
        "name": "Hero Section",
        "purpose": "Create an immediate, compelling introduction to Nirvaha's brand and core mission",
        "content": "Featured headline, subheadline, primary call-to-action button, background visuals with animated elements",
        "interaction": "Users can click the primary CTA to navigate to registration or explore further",
        "key_elements": "SpiritualOrb animated background, gradient overlays, prominent login/signup buttons"
    },
    {
        "name": "About Us Section",
        "purpose": "Communicate Nirvaha's philosophy, mission, and unique value proposition",
        "content": "Mission statement, brand story, core values, brief introduction to the holistic wellness approach",
        "interaction": "Informational section; users scroll to read content, may click related service links",
        "key_elements": "Text content, potentially decorative elements or sacred geometry patterns"
    },
    {
        "name": "Services Section",
        "purpose": "Present an overview of all major services and features available on the platform",
        "content": "Overview cards or sections for AI Chatbot, Meditation, Sound Healing, Community, Marketplace",
        "interaction": "Users can view service cards and click through to learn more or access specific features",
        "key_elements": "Service cards with icons, hover effects, navigation links to feature pages"
    },
    {
        "name": "Meditation & Wellness Features",
        "purpose": "Highlight the meditation platform's capabilities and wellness offerings",
        "content": "Preview of meditation types, guided sessions, wellness modules, certification offerings",
        "interaction": "Users can view feature previews and click to access the meditation platform",
        "key_elements": "Feature previews, MeditationPreview component, calls-to-action for meditation platform"
    },
    {
        "name": "Join Our Healing Community",
        "purpose": "Encourage users to become part of the Nirvaha community",
        "content": "Community benefits description, testimonials from community members, community features overview",
        "interaction": "Prominent CTA button directing users to community pages or registration",
        "key_elements": "CommunitySection component, membership benefits listing, registration buttons"
    },
    {
        "name": "User Feedback / Testimonials",
        "purpose": "Build credibility and trust through social proof and user experiences",
        "content": "User testimonials, success stories, rating/review information, user quotes",
        "interaction": "Users read testimonials; carousel or scroll through multiple testimonials",
        "key_elements": "Testimonial cards with user information, ratings, rotating carousel functionality"
    },
    {
        "name": "Gallery",
        "purpose": "Provide visual representation of platform features and the user experience",
        "content": "Screenshots or visual previews of key platform features, dashboard views, meditation sessions",
        "interaction": "Users browse visual content; may include lightbox or modal interactions",
        "key_elements": "Image gallery, possibly interactive carousel, hover effects"
    },
    {
        "name": "Collaborators / Partners",
        "purpose": "Display organizational partnerships and credibility markers",
        "content": "Partner logos, collaboration information, strategic partnership descriptions",
        "interaction": "Users view partner information; possible links to partner organizations",
        "key_elements": "Logo carousel or grid layout, partner descriptions"
    },
    {
        "name": "Contact Form",
        "purpose": "Enable direct communication between potential users and the Nirvaha team",
        "content": "Form fields for name, email, message, contact preferences",
        "interaction": "Users fill out the form and submit inquiries or feedback",
        "key_elements": "Input fields, validation, submit button, confirmation messaging"
    },
    {
        "name": "Footer",
        "purpose": "Provide navigation, legal compliance, and contact information",
        "content": "Quick links, social media links, contact information, legal pages, copyright information",
        "interaction": "Users access footer links for site navigation, legal information, or social media",
        "key_elements": "Multi-column layout, link sections, social media icons, contact details"
    }
]

for section in sections_breakdown:
    doc.add_heading(section["name"], level=2)
    doc.add_paragraph(f"Purpose: {section['purpose']}")
    doc.add_paragraph(f"Key Content: {section['content']}")
    doc.add_paragraph(f"User Interaction: {section['interaction']}")
    doc.add_paragraph(f"Key Elements: {section['key_elements']}")
    doc.add_paragraph()

# Section 5: Core Platform Features
doc.add_heading('5. Core Platform Features (Linked Pages / Modules)', level=1)

features = [
    {
        "name": "AI Spiritual Chatbot (ZenChat)",
        "description": "An intelligent conversational AI assistant designed to provide spiritual guidance, wellness advice, and personalized recommendations. The chatbot leverages artificial intelligence to understand user queries and provide thoughtful, contextual responses related to meditation, wellness, spirituality, and personal growth. Users can access ZenChat through a dedicated page and engage in continuous conversations for guidance and support."
    },
    {
        "name": "Meditation Platform (Mudra, Guided Meditation)",
        "description": "A comprehensive meditation and mindfulness platform offering various meditation techniques including Mudra-based practices, guided meditation sessions, breathing exercises, and wellness programs. The platform features categorized meditation content, progress tracking, personalized recommendations, and integration with the user's wellness dashboard."
    },
    {
        "name": "Sound Healing",
        "description": "A specialized module dedicated to sound-based wellness practices. Includes sound healing sessions, binaural beats, frequency-based healing audio, and related wellness content. Users can explore different sound healing modalities and integrate them into their wellness routine."
    },
    {
        "name": "Anonymous Discussion Rooms (Community)",
        "description": "Safe spaces within the community where users can engage in anonymous discussions on wellness topics, share experiences, seek advice, and support others without revealing their identity. This feature promotes open communication and community support within a secure environment."
    },
    {
        "name": "Companion Mode (Mentors & Guides)",
        "description": "A feature connecting users with experienced mentors, wellness guides, and spiritual teachers. Users can request guidance, book sessions, access mentorship programs, and receive personalized wellness recommendations from qualified companions within the Nirvaha ecosystem."
    },
    {
        "name": "Certification & Training Modules",
        "description": "Structured educational programs offering certifications in various wellness disciplines. Users can enroll in courses, complete modules, earn certifications, and become certified practitioners in areas such as meditation instruction, sound healing facilitation, and wellness coaching."
    },
    {
        "name": "Wellness Products Integration (Marketplace)",
        "description": "An integrated marketplace for wellness products including meditation aids, sound healing devices, wellness supplements, books, and related items. Users can browse products, make purchases, and have items delivered directly to them, enhancing their wellness journey."
    },
    {
        "name": "User Dashboard",
        "description": "A personalized dashboard providing users with an overview of their wellness journey. The dashboard displays progress metrics, meditation statistics, upcoming sessions, community activity, marketplace orders, and personalized recommendations based on user behavior and preferences."
    }
]

for feature in features:
    doc.add_heading(feature["name"], level=2)
    doc.add_paragraph(feature["description"])
    doc.add_paragraph()

# Section 6: Navigation Structure
doc.add_heading('6. Navigation Structure', level=1)

nav_heading = doc.add_heading('Header Navigation', level=2)
header_nav = """The main navigation header is accessible from all pages and provides quick access to primary platform sections. The header typically includes:
• Logo / Branding (links to home)
• Service Links (Meditation, AI Chatbot, Sound Healing, Community, Marketplace)
• Account Links (Dashboard, Profile, Logout for authenticated users)
• Primary CTA Buttons (Sign Up, Get Started, Login)
• Mobile Menu (Hamburger menu for responsive design on smaller devices)"""
doc.add_paragraph(header_nav)
doc.add_paragraph()

cta_heading = doc.add_heading('Call-to-Action Buttons and Navigation', level=2)
cta_text = """CTA buttons are strategically placed throughout the website to guide users toward key conversion points:
• Primary CTA: "Get Started" or "Join Now" - directs users to registration
• Feature CTAs: Service-specific buttons directing to individual feature pages
• Community CTA: Encourages participation in community features
• Login / Signup: Authentication buttons accessible from multiple locations
• Explore Buttons: Encourage users to learn more about specific features"""
doc.add_paragraph(cta_text)
doc.add_paragraph()

auth_heading = doc.add_heading('Authentication Flow', level=2)
auth_text = """The authentication system includes:
• Login Page: For returning users to authenticate with email/password or social login options
• Registration Page: For new users to create accounts with email, password, and profile information
• Profile Page: For users to manage account settings, preferences, and personal information
• Logout: Available in the user menu to securely end the user session
• Password Recovery: Option to reset forgotten passwords via email"""
doc.add_paragraph(auth_text)
doc.add_paragraph()

# Section 7: Footer Structure
doc.add_heading('7. Footer Structure', level=1)

footer_text = """The footer appears at the bottom of every page and provides essential navigation and information:"""
doc.add_paragraph(footer_text)
doc.add_paragraph()

footer_sections = [
    {
        "category": "Quick Links",
        "items": "Home, About, Services, Blog, FAQ, Contact"
    },
    {
        "category": "Features",
        "items": "Meditation Platform, AI Chatbot, Sound Healing, Community, Marketplace, Certification"
    },
    {
        "category": "Account",
        "items": "Dashboard, Profile, My Orders, My Sessions, Preferences"
    },
    {
        "category": "Legal & Compliance",
        "items": "Privacy Policy, Terms of Service, Cookie Policy, Disclaimer, GDPR Compliance"
    },
    {
        "category": "Contact Information",
        "items": "Email address, Phone number, Contact form link, Support portal link"
    },
    {
        "category": "Social Media",
        "items": "Links to Facebook, Instagram, Twitter, LinkedIn, YouTube"
    },
    {
        "category": "Newsletter",
        "items": "Email subscription form for wellness tips and platform updates"
    },
    {
        "category": "Copyright",
        "items": "Copyright notice, year, company name, all rights reserved statement"
    }
]

for section in footer_sections:
    doc.add_paragraph(f"{section['category']}: {section['items']}", style='List Bullet')

doc.add_paragraph()

# Closing section
doc.add_heading('Document Summary', level=1)
closing_text = """This document provides a comprehensive structural overview of the Nirvaha website. The platform is designed to serve as a holistic wellness ecosystem, combining AI-powered guidance, meditation practices, sound healing, community support, and marketplace integration. The structure emphasizes user engagement, ease of navigation, and continuous accessibility to core features across all pages and devices.

The website prioritizes user experience through clear information hierarchy, strategic call-to-action placement, and seamless navigation between the landing page and specialized feature pages. All sections work cohesively to guide users from initial awareness through active engagement with the platform's services and community.

For questions or clarifications regarding this document, please contact the product or development team."""
doc.add_paragraph(closing_text)

# Save the document
output_path = r'c:\Preetham\B.TECH\INTERN\NIRVAHA\AKSHU\nirvaha-frontend\Dashboard\Nirvaha_Website_Structure.docx'
doc.save(output_path)
print(f"Document successfully created: {output_path}")
